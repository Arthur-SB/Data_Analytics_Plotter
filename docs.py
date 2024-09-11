from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# dict_docs = {
# 'imagem 1': {'title': 'Temperatura do ar por log do teste', 'font': 'Acervo Pessoal', 'description': 'Texto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemplo', 'filename': 'E:/GitHub/Data_Analytics_Plotter/plotter/logs/logs X Air Temp Inlet C.png'},
# 'imagem 2': {'title': 'Pressão de freio da dianteira por log do teste', 'font': 'Acervo Pessoal', 'description': 'Texto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemploTexto exemplo', 'filename': 'E:/GitHub/Data_Analytics_Plotter/plotter/logs/logs X Alias Brake Pres Front MPa.png'},
# 'name_report': 'Validação da capacidade de frenagem do protótipo de freio do e21',
# 'objetivo': 'Através de testes como pressão máxima das linhas de freio, variação do bias e hysteres do freio e o quanto isso impacta no wheel locked',
# 'responsible': 'Congonhas e Spam',
# 'driver': 'Rorato e Rubinho',
# 'date': '12 de dezembro',
# 'hour_in': '8:23',
# 'hour_out': '18:17',
# 'local': 'Bebedouro',
# 'problem': '	Desgaste Prematuro das Pastilhas de Freio Durante as primeiras rodadas de frenagem brusca a 60 km/h, observou-se um desgaste anormal das pastilhas de freio. Isso causou um aumento gradual na distância de frenagem, comprometendo o desempenho do sistema. Após uma inspeção visual, a equipe constatou que o material das pastilhas não suportava as altas temperaturas geradas durante a frenagem intensa, perdendo eficiência rapidamente. Foi decidido substituir as pastilhas por um material com maior resistência térmica, feito de composto cerâmico-carbono. Com isso, os testes subsequentes mostraram que o desgaste diminuiu significativamente e a eficiência da frenagem se estabilizou. Superaquecimento dos Discos de Freio Após várias frenagens consecutivas, os discos de freio atingiram temperaturas excessivas, o que resultou em um fenômeno conhecido como "fading" (perda temporária da capacidade de frenagem). Isso aumentou a distância de parada e gerou riscos ao piloto. A equipe identificou que a ventilação dos discos de freio era inadequada para dissipar o calor acumulado, especialmente durante frenagens repetidas. Um sistema de ventilação forçada foi instalado nos discos, aumentando a circulação de ar e melhorando a dissipação de calor. Além disso, a equipe alterou o cronograma de testes, incluindo pausas para permitir o resfriamento dos componentes entre as sessões de frenagem. Após essas modificações, o problema de superaquecimento foi resolvido.',
# 'conclusion': 'Após os ajustes e correções implementados durante o teste, o sistema de freios passou a atender aos requisitos de desempenho definidos pela equipe. O protótipo demonstrou alta capacidade de frenagem, tanto em superfícies secas quanto molhadas, com distâncias de parada dentro dos limites esperados. O teste foi considerado um sucesso, e o protótipo de freio foi aprovado para ser utilizado na competição.'
# }

def report(dict_docs):

    general = f"O teste de foi realizado no dia " + dict_docs["date"] + ", no local " + dict_docs["local"] + ". As atividades começaram às " + dict_docs["hour_in"] + " e foram concluídas às " + dict_docs["hour_out"] + ". O responsável pelo teste foi " + dict_docs["responsible"] + ", com os pilotos " + dict_docs["driver"] + " conduzindo o carro durante as fases do experimento."


    current_year = datetime.now().strftime("%Y")
    current_month = datetime.now().strftime("%m")
    current_day = datetime.now().strftime("%d")
    i=1

    # Criando o documento
    doc = Document()

    # Ajustar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(85)  # 3 cm
        section.bottom_margin = Pt(57)  # 2 cm
        section.left_margin = Pt(85)  # 3 cm
        section.right_margin = Pt(57)  # 2 cm
    
#--------------- Cover ---------------#

    header = doc.add_paragraph()
    run = header.add_run("Universidade de São Paulo\nEscola de Engenharia de São Carlos\nEESC-USP Formula SAE")
    run.font.size = Pt(14)
    run.font.name = 'Segoe UI'
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph("\n" * 3)

    title = doc.add_paragraph()
    run = title.add_run(dict_docs["name_report"])
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = 'Segoe UI'
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph("\n" * 3)

    logo = doc.add_paragraph()
    run = logo.add_run()
    run.add_picture('resource/for_preto.png') 
    logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph("\n" * 10)
    
    footer = doc.add_paragraph()
    run = footer.add_run(f"SÃO CARLOS - SP\n{current_year}")
    run.font.size = Pt(12)
    run.font.name = 'Segoe UI'
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

#--------------- Geral ---------------#
    title = doc.add_heading(level=1)
    run = title.add_run("1. Geral")

    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.all_caps = True
    run.bold = True
    run.font.size = Pt(14)

    title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph_format = title.paragraph_format
    paragraph_format.space_before = Pt(24)
    paragraph_format.line_spacing = 1.5


    paragraph = doc.add_paragraph()
    run = paragraph.add_run(general)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(12)

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5

#--------------- Objetivo ---------------#
    title = doc.add_heading(level=1)
    run = title.add_run("2. Objetivo")

    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.all_caps = True
    run.bold = True
    run.font.size = Pt(14)

    title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph_format = title.paragraph_format
    paragraph_format.space_before = Pt(24)
    paragraph_format.line_spacing = 1.5


    paragraph = doc.add_paragraph()
    run = paragraph.add_run(dict_docs['objetivo'])
    run.font.name = 'Segoe UI'
    run.font.size = Pt(12)

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5

    doc.add_page_break()


#--------------- Logs ---------------#
    title = doc.add_heading(level=1)
    run = title.add_run("3. Logs")

    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.all_caps = True
    run.bold = True
    run.font.size = Pt(14)

    title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph_format = title.paragraph_format
    paragraph_format.space_before = Pt(24)
    paragraph_format.line_spacing = 1.5

    for key in dict_docs:
        if key.startswith('imagem'):
            value = dict_docs[key]

            # Adiciona e formata o título
            title_paragraph = doc.add_paragraph()
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title_paragraph.add_run(f"Imagem {i} - {value['title']}")
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Segoe UI'        
            run.font.color.rgb = RGBColor(0, 0, 0)

            image_paragraph = doc.add_paragraph()
            image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = image_paragraph.add_run()
            run.add_picture(value['filename'], width=Inches(6.5))
            
            # Adiciona e formata a fonte
            font_paragraph = doc.add_paragraph()
            font_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            paragraph_format.space_before = Pt(24)
            run = font_paragraph.add_run(f"Fonte: {value['font']}")
            run.font.size = Pt(10)
            run.font.name = 'Segoe UI'    
            run.font.color.rgb = RGBColor(0, 0, 0)

            
            # Adiciona e formata a descrição
            description_paragraph = doc.add_paragraph()
            description_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            run = description_paragraph.add_run(value['description'])
            run.font.size = Pt(12)
            run.font.name = 'Segoe UI'    
            run.font.color.rgb = RGBColor(0, 0, 0)

            i += 1
            doc.add_page_break()
    
    


#--------------- Problemas e soluções ---------------#
    title = doc.add_heading(level=1)
    run = title.add_run('4. Problemas e Soluções')

    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.all_caps = True
    run.bold = True
    run.font.size = Pt(14)

    title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph_format = title.paragraph_format
    paragraph_format.space_before = Pt(24)
    paragraph_format.line_spacing = 1.5


    paragraph = doc.add_paragraph()
    run = paragraph.add_run(dict_docs['problem'])
    run.font.name = 'Segoe UI'
    run.font.size = Pt(12)

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5

#--------------- Conclusão ---------------#
    title = doc.add_heading(level=1)
    run = title.add_run('5. Conclusão')

    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.all_caps = True
    run.bold = True
    run.font.size = Pt(14)

    title.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph_format = title.paragraph_format
    paragraph_format.space_before = Pt(24)
    paragraph_format.line_spacing = 1.5


    paragraph = doc.add_paragraph()
    run = paragraph.add_run(dict_docs['conclusion'])
    run.font.name = 'Segoe UI'
    run.font.size = Pt(12)

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = 1.5

#--------------- Salvar Arquivo ---------------#
    doc.save(f'Test_Report_{current_year}_{current_month}_{current_day}.docx')

# report(dict_docs)